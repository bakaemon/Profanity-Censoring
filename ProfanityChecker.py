import os
from docx import Document

from data.Loader import Loader


def get_ext(filepath):
    return os.path.splitext(filepath)[-1].lower()


def check_ext(filepath):
    ext = get_ext(filepath)
    return ext in [".txt", ".doc", ".docx"]


class ProfanityChecker:
    _badwords = []
    _document = []

    def load_badwords_from_file(self, file_path: str):
        with Loader("Retrieving data...", end="Retrieving data...Done!"):
            with open(file_path, "r") as file:
                for line in file:
                    self._badwords.append(line.replace("\n", ""))
                file.close()
        return self

    def load_document(self, doc_file: str):
        if check_ext(doc_file):
            if get_ext(doc_file) in [".doc", ".docx"]:
                loader = Loader("Reading document...", end="Reading document...Done!")
                doc_file = Document(doc_file)
                loader.stop()
                with Loader("Saving to temp...", end="Saving to temp...Done!"):
                    for p in doc_file.paragraphs:
                        self._document.append(p.text)
            elif get_ext(doc_file) == ".txt":
                with Loader("Saving to temp...", end="Saving to temp...Done!"):
                    with open(doc_file, "r") as file:
                        for line in file:
                            self._document.append(line.replace("\n", ""))
                        file.close()

        return self

    def censor(self, censor_char="*"):
        with Loader("Processing paragraphs...", end="Processing paragraphs...Done!"):
            for index, paragraph in enumerate(self._document):
                self._document[index] = self.wfw_censor(paragraph, censor_char)

        return self

    def wfw_censor(self, sentence, censor_char="*"):
        """
        Standalone censor function word for word (wfw).
        """
        sentence = sentence.split()
        for index, word in enumerate(sentence):
            if any(badword in word for badword in self._badwords):
                sentence[index] = "".join([censor_char if c.isalpha() else c for c in word])

        return " ".join(sentence)

    def export(self, new_document: str):
        ext = os.path.splitext(new_document)[-1].lower()
        if ext not in [".txt", ".doc", ".docx"]:
            print("Extension is not supported.")
            return
        with Loader("Exporting in progress...", end="Exporting in progress...Done!"):
            if ext == ".txt":
                with open(new_document, "w") as file:
                    for p in self._document:
                        file.write(p + "\n")
                    file.close()
            elif ext in [".doc", ".docx"]:
                newfile = Document()
                for p in self._document:
                    newfile.add_paragraph(p)
                newfile.save(new_document)
